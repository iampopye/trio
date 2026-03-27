"""File upload handler -- extract text content from uploaded documents.

Supports: PDF, DOCX, XLSX, CSV, TXT, Python, JS, JSON, Markdown,
          HTML, images (OCR placeholder), and more.
"""

import csv
import io
import json
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

# Max file size: 25MB
MAX_FILE_SIZE = 25 * 1024 * 1024

SUPPORTED_EXTENSIONS = {
    # Documents
    ".pdf", ".docx", ".xlsx", ".csv", ".txt", ".md", ".rtf",
    # Code
    ".py", ".js", ".ts", ".jsx", ".tsx", ".java", ".go", ".rs",
    ".rb", ".php", ".c", ".cpp", ".h", ".cs", ".swift", ".kt",
    ".sh", ".bash", ".zsh", ".ps1", ".bat",
    # Config/Data
    ".json", ".yaml", ".yml", ".toml", ".xml", ".ini", ".env",
    ".sql", ".graphql",
    # Web
    ".html", ".htm", ".css", ".scss", ".less", ".svg",
    # Images (basic info)
    ".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp",
    # Other
    ".log", ".diff", ".patch",
}

TEXT_EXTENSIONS = {
    ".txt", ".md", ".py", ".js", ".ts", ".jsx", ".tsx", ".java",
    ".go", ".rs", ".rb", ".php", ".c", ".cpp", ".h", ".cs",
    ".swift", ".kt", ".sh", ".bash", ".zsh", ".ps1", ".bat",
    ".json", ".yaml", ".yml", ".toml", ".xml", ".ini", ".env",
    ".sql", ".graphql", ".html", ".htm", ".css", ".scss",
    ".less", ".svg", ".log", ".diff", ".patch", ".rtf",
}


def extract_text(file_bytes: bytes, filename: str) -> dict:
    """Extract text content from a file.

    Returns:
        {
            "filename": str,
            "type": str,
            "size": int,
            "content": str,       # extracted text
            "preview": str,       # first 200 chars
            "pages": int | None,  # for PDFs
            "error": str | None,
        }
    """
    ext = Path(filename).suffix.lower()
    size = len(file_bytes)

    result = {
        "filename": filename,
        "type": ext.lstrip("."),
        "size": size,
        "content": "",
        "preview": "",
        "pages": None,
        "error": None,
    }

    if size > MAX_FILE_SIZE:
        result["error"] = f"File too large ({size // (1024*1024)}MB). Max 25MB."
        return result

    if ext not in SUPPORTED_EXTENSIONS:
        result["error"] = f"Unsupported file type: {ext}"
        return result

    try:
        if ext == ".pdf":
            result["content"], result["pages"] = _extract_pdf(file_bytes)
        elif ext == ".docx":
            result["content"] = _extract_docx(file_bytes)
        elif ext == ".xlsx":
            result["content"] = _extract_xlsx(file_bytes)
        elif ext == ".csv":
            result["content"] = _extract_csv(file_bytes)
        elif ext in (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp"):
            result["content"] = _describe_image(file_bytes, filename)
            result["type"] = "image"
        elif ext in TEXT_EXTENSIONS:
            result["content"] = _extract_text_file(file_bytes)
        else:
            result["content"] = _extract_text_file(file_bytes)
    except Exception as e:
        logger.error(f"Error extracting {filename}: {e}")
        result["error"] = str(e)

    # Truncate very long content
    if len(result["content"]) > 50000:
        result["content"] = result["content"][:50000] + "\n\n... (truncated, file too long)"

    result["preview"] = result["content"][:200] if result["content"] else ""
    return result


def _extract_pdf(data: bytes) -> tuple[str, int]:
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        return "Error: PyPDF2 not installed. Run: pip install PyPDF2", 0

    reader = PdfReader(io.BytesIO(data))
    pages = len(reader.pages)
    text_parts = []

    for i, page in enumerate(reader.pages):
        page_text = page.extract_text() or ""
        if page_text.strip():
            text_parts.append(f"--- Page {i+1} ---\n{page_text}")

    return "\n\n".join(text_parts) or "(No extractable text in PDF)", pages


def _extract_docx(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "Error: python-docx not installed. Run: pip install python-docx"

    doc = Document(io.BytesIO(data))
    parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Also extract tables
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            parts.append("\n".join(rows))

    return "\n\n".join(parts) or "(Empty document)"


def _extract_xlsx(data: bytes) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "Error: openpyxl not installed. Run: pip install openpyxl"

    wb = load_workbook(io.BytesIO(data), read_only=True)
    parts = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(max_row=200, values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            parts.append(f"--- Sheet: {sheet_name} ---\n" + "\n".join(rows))

    wb.close()
    return "\n\n".join(parts) or "(Empty spreadsheet)"


def _extract_csv(data: bytes) -> str:
    text = data.decode("utf-8", errors="replace")
    reader = csv.reader(io.StringIO(text))
    rows = []
    for i, row in enumerate(reader):
        if i > 500:
            rows.append("... (500+ rows, truncated)")
            break
        rows.append(" | ".join(row))
    return "\n".join(rows) or "(Empty CSV)"


def _extract_text_file(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, Exception):
            continue
    return data.decode("utf-8", errors="replace")


def _describe_image(data: bytes, filename: str) -> str:
    """Generate a text description of an image."""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data))
        w, h = img.size
        mode = img.mode
        fmt = img.format or Path(filename).suffix.upper()
        return (
            f"[Image: {filename}]\n"
            f"Dimensions: {w}x{h} pixels\n"
            f"Format: {fmt}, Mode: {mode}\n"
            f"Size: {len(data) // 1024}KB\n\n"
            f"(Image content analysis requires a vision-capable model. "
            f"The image has been received and saved.)"
        )
    except Exception:
        return f"[Image: {filename}] ({len(data) // 1024}KB)"
